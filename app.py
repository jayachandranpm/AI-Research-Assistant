# --- START OF FILE app.py ---

import os
import uuid
import time
import logging
import random # Keep for potential future use (jitter)
import html
import re
import traceback
from io import BytesIO

from flask import (
    Flask, render_template, request, jsonify,
    send_file, make_response, current_app
)
from dotenv import load_dotenv
import google.generativeai as genai
from duckduckgo_search import DDGS
import requests
from bs4 import BeautifulSoup
import trafilatura
import mistune

from docx import Document
from docx.shared import Pt, Inches
from xhtml2pdf import pisa

# --- Configuration ---
load_dotenv()
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - [%(funcName)s:%(lineno)d] - %(message)s'
)

# --- Constants ---
QUICK_SEARCH_RESULTS = 7
DEEP_SEARCH_RESULTS = 18
SCRAPE_TIMEOUT = 8
MAX_CONTENT_LENGTH_PER_SITE = 15000
MAX_TOTAL_CONTENT_LENGTH = 200000
REPORT_STORE_MAX_ITEMS = 100
SOURCE_PREVIEW_LENGTH = 300

# --- In-memory storage ---
report_store = {}
report_order = []

def add_to_report_store(report_id, data):
    """
    Add a report to the in-memory store, managing maximum number of items.

    Args:
        report_id (str): Unique identifier for the report
        data (dict): Report data to store
    """
    if len(report_store) >= REPORT_STORE_MAX_ITEMS:
        try:
            oldest_id = report_order.pop(0)
            if oldest_id in report_store:
                del report_store[oldest_id]
                logging.info(f"Removed oldest report ({oldest_id}).")
        except IndexError:
            pass  # No reports to remove

    report_store[report_id] = data
    report_order.append(report_id)

# --- Helper Functions ---

def perform_search(query: str, num_results: int) -> list[dict]:
    """
    Perform a web search using DuckDuckGo.

    Args:
        query (str): Search query
        num_results (int): Maximum number of results to return

    Returns:
        list[dict]: List of search results with title, URL, and body
    """
    logging.info(f"Searching for '{query}' (max {num_results} results).")
    results = []
    try:
        with DDGS(timeout=15) as ddgs:
            search_results = list(ddgs.text(query, max_results=num_results))

        if search_results:
            results = [
                {
                    "title": r.get("title", ""),
                    "url": r.get("href", ""),
                    "body": r.get("body", "")
                }
                for r in search_results if r.get("href")
            ]
    except Exception as e:
        logging.error(f"DDG search failed: {e}", exc_info=True)

    logging.info(f"Found {len(results)} results.")
    return results

def scrape_url(url: str, timeout: int = SCRAPE_TIMEOUT) -> str | None:
    """
    Scrape the main content from a given URL, with Trafilatura fix.

    Args:
        url (str): URL to scrape
        timeout (int): Request timeout in seconds

    Returns:
        str | None: Scraped content or None if scraping fails
    """
    logging.info(f"Attempting scrape: {url}")
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.0.0 Safari/537.36 ResearchAssistantBot/1.0'
    }

    # 1. Try Trafilatura
    try:
        downloaded = trafilatura.fetch_url(url) # Removed timeout kwarg
        if downloaded:
            # --- FIX: Use 'txt' for output format ---
            main_content = trafilatura.extract(
                downloaded, include_comments=False,
                include_tables=False, output_format='txt' # CORRECTED format
            )
            # --- End FIX ---
            if main_content and len(main_content) > 100:
                logging.info(f"Trafilatura success ({len(main_content)} chars): {url}")
                return main_content.strip()[:MAX_CONTENT_LENGTH_PER_SITE]
            else:
                logging.warning(f"Trafilatura extracted minimal/no content (len={len(main_content or '')}) from: {url}")
        else:
            logging.warning(f"Trafilatura could not fetch URL (returned None): {url}")
    except TypeError as te:
         # Catch TypeErrors which might indicate wrong args for the installed version
         logging.error(f"Trafilatura extract/fetch failed for {url}: {te.__class__.__name__} - {te}", exc_info=False)
    except Exception as e:
        logging.error(f"Trafilatura failed for {url}: {e.__class__.__name__} - {e}", exc_info=False)

    # 2. Fallback to BeautifulSoup
    logging.info(f"Falling back to BeautifulSoup for: {url}")
    try:
        response = requests.get(url, timeout=timeout, headers=headers, verify=True)
        response.raise_for_status()
        content_type = response.headers.get('Content-Type', '').lower()
        if 'html' not in content_type: logging.warning(f"Skipping non-HTML ({content_type}): {url}"); return None
        if len(response.content) > 7_000_000: logging.warning(f"Content size exceeds limit: {url}"); return None

        soup = BeautifulSoup(response.content, 'lxml')
        potential_containers = [soup.find('article'), soup.find('main'), soup.find('div', id='content'), soup.find('div', class_='content'), soup.find('div', id='main-content'), soup.find('div', class_='main-content'), soup.find('div', class_='entry-content'), soup.find('div', role='main'), soup]
        text = ""
        for container in potential_containers:
            if container:
                paragraphs = container.find_all('p', recursive=True)
                if paragraphs:
                    extracted_text = ' '.join(p.get_text(strip=True) for p in paragraphs)
                    if len(extracted_text) > 200: text = extracted_text; logging.info(f"BS4 using container '{container.name}' ({len(text)} chars): {url}"); break
        if not text:
             all_paragraphs = soup.find_all('p')
             if all_paragraphs:
                  extracted_text = ' '.join(p.get_text(strip=True) for p in all_paragraphs)
                  if len(extracted_text) > 100: text = extracted_text; logging.info(f"BS4 joining all 'p' tags ({len(text)} chars): {url}");
        if text: return text.strip()[:MAX_CONTENT_LENGTH_PER_SITE]
        else: body_snippet = soup.body.get_text(strip=True, separator=' ')[0:500] if soup.body else "No body"; logging.warning(f"BS4 fallback found no content. Snippet: '{body_snippet}...': {url}"); return None
    except requests.exceptions.Timeout: logging.error(f"Timeout (Requests): {url}"); return None
    except requests.exceptions.HTTPError as e: logging.error(f"HTTP Error {e.response.status_code} (Requests): {url}"); return None
    except requests.exceptions.RequestException as e: logging.error(f"RequestException (Requests) {url}: {e}", exc_info=False); return None
    except Exception as e: logging.error(f"Generic Error BS4 fallback {url}:", exc_info=True); return None


def scrape_urls(urls: list[str]) -> list[dict]:
    """Scrapes multiple URLs."""
    scraped_data = []; total_content_length = 0; urls_attempted = set()
    logging.info(f"Starting scrape for {len(urls)} URLs.")
    for i, url in enumerate(urls):
        if url in urls_attempted: continue; urls_attempted.add(url)
        if total_content_length >= MAX_TOTAL_CONTENT_LENGTH: logging.warning(f"Reached max total content length."); break
        content = scrape_url(url)
        if content: scraped_data.append({"id": len(scraped_data), "url": url, "text": content}); total_content_length += len(content); logging.debug(f"Scrape success {url}")
        time.sleep(0.3) # Polite delay
    logging.info(f"Finished scraping. Success: {len(scraped_data)}/{len(urls_attempted)} URLs. Total length: {total_content_length}")
    return scraped_data


def format_context_for_llm(scraped_data: list[dict]) -> str:
    """Formats scraped data for LLM context."""
    context_str = ""
    for item in scraped_data: context_str += f"Source [{item['id']+1}] ({item['url']}):\n{item['text']}\n\n---\n\n"
    return context_str


def preprocess_llm_text_for_citations(text: str) -> str:
    """
    Splits combined citations e.g., [1, 2] -> [1][2] and ensures proper spacing
    between consecutive citation markers.
    """
    # First handle comma-separated citations
    def replacer(match):
        numbers = match.group(1).split(',')
        return " ".join(f"[{num.strip()}]" for num in numbers if num.strip().isdigit())
    
    # Replace comma-separated citations
    processed_text = re.sub(r"\[\s*(\d+\s*(?:,\s*\d+\s*)*)\s*\]", replacer, text)
    
    # Add space between consecutive citation markers to ensure they're processed separately
    processed_text = re.sub(r"(\[\d+\])(\[\d+\])", r"\1 \2", processed_text)
    
    return processed_text


def process_citation_markers_in_html(html_content: str) -> str:
    """
    Improved citation marker processing for HTML content.
    Handles consecutive citations properly with spacing.
    """
    # First pass: Replace all citation markers
    def replace_citation_html(match):
        num = int(match.group(1))
        return f"<sup><a href='#' class='citation-marker' data-citation-index='{num-1}' aria-label='Citation {num}'>[{num}]</a></sup>"

    processed_html = re.sub(r"(?<![!\]/a-zA-Z0-9])\[(\d+)\](?!\])", replace_citation_html, html_content)

    # Second pass: Add spacing between consecutive sup tags for better rendering
    processed_html = re.sub(r'(</sup>)(<sup>)', r'\1 \2', processed_html)

    return processed_html


def generate_chunked_deep_research(model, prompt, safety_settings, query, context_string):
    """
    Generates deep research content in chunks to avoid token limit issues.
    Returns the combined text from all chunks.
    """
    logging.info("Using chunked generation for deep research")

    # First, generate the structure (title, abstract, intro)
    structure_prompt = f"""
    User Query: "{query}"

    Based on the provided sources, generate ONLY the Title, Abstract, and Introduction sections
    for a comprehensive academic research article. Focus on creating a strong foundation for the
    article with a clear research question, context, and objectives.

    Sources:
    --- START OF SOURCES ---
    {context_string}
    --- END OF SOURCES ---
    """

    structure_response = model.generate_content(
        structure_prompt,
        safety_settings=safety_settings,
        generation_config=genai.types.GenerationConfig(max_output_tokens=2000, temperature=0.6)
    )

    if not structure_response.candidates or not hasattr(structure_response.candidates[0].content, 'parts'):
        logging.error("Failed to generate article structure")
        return None

    # Get the initial content
    generated_text = structure_response.text

    # Now generate the literature review and analysis
    middle_prompt = f"""
    User Query: "{query}"

    Based on the provided sources, generate ONLY the Literature Review and Analysis/Discussion sections
    for a comprehensive academic research article. The Title, Abstract, and Introduction have already been generated.
    Focus on thorough analysis of existing research and detailed discussion of findings.

    Sources:
    --- START OF SOURCES ---
    {context_string}
    --- END OF SOURCES ---
    """

    middle_response = model.generate_content(
        middle_prompt,
        safety_settings=safety_settings,
        generation_config=genai.types.GenerationConfig(max_output_tokens=3000, temperature=0.6)
    )

    if middle_response.candidates and hasattr(middle_response.candidates[0].content, 'parts'):
        generated_text += "\n\n" + middle_response.text

    # Finally, generate the conclusion and references
    conclusion_prompt = f"""
    User Query: "{query}"

    Based on the provided sources, generate ONLY the Conclusion and References sections
    for a comprehensive academic research article. The previous sections have already been generated.
    Focus on summarizing key findings and properly formatting all references.

    Sources:
    --- START OF SOURCES ---
    {context_string}
    --- END OF SOURCES ---
    """

    conclusion_response = model.generate_content(
        conclusion_prompt,
        safety_settings=safety_settings,
        generation_config=genai.types.GenerationConfig(max_output_tokens=2000, temperature=0.6)
    )

    if conclusion_response.candidates and hasattr(conclusion_response.candidates[0].content, 'parts'):
        generated_text += "\n\n" + conclusion_response.text

    return generated_text


def create_gemini_prompt(query: str, context_string: str, depth: str) -> str:
    """Creates the prompt, adjusted for depth and emphasizing deep research requirements."""
    base_instructions = f"""User Query: "{query}"

Sources:
--- START OF SOURCES ---
{context_string}
--- END OF SOURCES ---

General Instructions:
1.  Analyze the User Query and ALL provided Sources meticulously. Ignore irrelevant sources.
2.  **Cite (Inline):** Add `[N]` after information from Source N. Use individual markers `[1][4][5]` for combined info. Place before punctuation. **Accuracy is critical.**
3.  **Source Reliance:** Base the response *exclusively* on the provided sources. NO outside knowledge.
4.  **Clarity & Structure:** Use clear language and Markdown formatting.
5.  **Handling Gaps:** If sources lack info, state it clearly. Do not invent.
6.  **Tone:** Objective, factual, neutral.
7.  **Output:** Generate ONLY the Markdown answer, starting directly without preamble.
"""
    if depth == 'deep':
        deep_specific_instructions = f"""
COMPREHENSIVE ACADEMIC RESEARCH ARTICLE GENERATION GUIDELINES

Research Query: "{query}"

ARTICLE STRUCTURE AND EXPECTATIONS:
1. Length: Generate a comprehensive research article of approximately 10-15 pages (2500-3750 words)
2. Academic Rigor: Produce a scholarly, well-researched, and critically analyzed response
3. Formatting: Use academic writing conventions with clear structure

DETAILED ARTICLE COMPONENTS:

I. Title Page
- Concise, informative title reflecting the research query
- Subtitle if appropriate
- Implied academic affiliation (e.g., "Interdisciplinary Research Institute")

II. Abstract (250-300 words)
- Clearly state research query
- Summarize key findings
- Highlight research significance
- Briefly mention methodological approach

III. Introduction (500-600 words)
- Contextualize the research query
- State research objectives
- Explain significance and relevance
- Provide theoretical or conceptual framework
- Present central research questions or hypotheses

IV. Comprehensive Literature Review (800-1000 words)
- Systematic review of existing research
- Critical analysis of current knowledge
- Identify gaps and controversies
- Synthesize perspectives from multiple sources
- Demonstrate deep understanding of the topic

V. Detailed Analysis and Discussion (1200-1500 words)
- In-depth exploration of research findings
- Integrate evidence from provided sources
- Present multiple perspectives
- Critical interpretation of information
- Address potential counterarguments
- Connect findings to broader context

VI. Conclusion (400-500 words)
- Summarize key findings
- Reflect on research implications
- Suggest future research directions
- Provide definitive insights

VII. References
- Comprehensive list of sources
- Formatted in APA, MLA, or Chicago style
- Include all cited sources from web scraping
- Ensure proper academic citation format

SPECIFIC RESEARCH GUIDELINES:
- Use academic, formal language
- Maintain objective, analytical tone
- Critically evaluate and synthesize information
- Provide nuanced, balanced perspectives
- Demonstrate scholarly depth and intellectual rigor

CITATION INSTRUCTIONS:
- Use numbered citation markers [1], [2], etc.
- Cite sources for all claims and data points
- Paraphrase and quote appropriately
- Maintain academic integrity

CONTEXT FROM WEB SOURCES:
{context_string}

FINAL INSTRUCTIONS:
- Prioritize accuracy and scholarly integrity
- Do not fabricate information
- If sources are insufficient, clearly state limitations
- Aim for a comprehensive, authoritative response

Generate the complete research article following these guidelines."""
        return base_instructions + deep_specific_instructions + "\nSynthesized Deep Research Article (Markdown format):"
    else: # Quick depth
        quick_specific_instructions = """
Specific Instructions for Quick Answer:
*   **Goal:** Concise, informative answer synthesizing key points from sources.
*   **Structure:** Use paragraphs, `### Subheadings` (optional), `* Bullet points`.
"""
        return base_instructions + quick_specific_instructions + "\nSynthesized Answer (Markdown format):"


def synthesize_with_gemini(query: str, scraped_data: list[dict], api_key: str, depth: str) -> dict | None:
    """Synthesizes content using Gemini, adds spacing for citations."""
    if not api_key: return {"error": "AI key missing."}
    if not scraped_data: return {"error": "No content to synthesize."}

    genai.configure(api_key=api_key)
    context_string = format_context_for_llm(scraped_data)
    prompt = create_gemini_prompt(query, context_string, depth)

    estimated_tokens = len(prompt) / 4
    logging.info(f"Sending prompt to Gemini for '{depth}'. Estimated context: ~{estimated_tokens:.0f} tokens.")

    model_name = 'gemini-2.0-flash' if depth == 'deep' else 'gemini-2.0-flash'
    max_tokens = 8192 # Keep generous
    logging.info(f"Using model: {model_name} (max_tokens={max_tokens})")
    model = genai.GenerativeModel(model_name)
    safety_settings = [ {"category": c, "threshold": "BLOCK_MEDIUM_AND_ABOVE"} for c in ["HARM_CATEGORY_HARASSMENT", "HARM_CATEGORY_HATE_SPEECH", "HARM_CATEGORY_SEXUALLY_EXPLICIT", "HARM_CATEGORY_DANGEROUS_CONTENT"] ]

    try:
        # For deep research, use chunked generation to avoid token limit issues
        if depth == 'deep':
            generated_text_raw = generate_chunked_deep_research(model, prompt, safety_settings, query, context_string)
            if not generated_text_raw:
                return {"error": "Failed to generate deep research content"}
        else:
            # For quick research, use single generation
            response = model.generate_content(prompt, safety_settings=safety_settings, generation_config=genai.types.GenerationConfig(max_output_tokens=max_tokens, temperature=0.6))

            # Check for errors in response
            if response.candidates and response.candidates[0].finish_reason.name != "STOP": reason = response.candidates[0].finish_reason.name; logging.warning(f"Gemini finished unexpectedly: {reason}"); error_map = {"MAX_TOKENS": "AI response cut short.", "SAFETY": "Blocked by safety.", "RECITATION": "Blocked (recitation)."}; return {"error": error_map.get(reason, f"AI failed (Reason: {reason}).")}
            if not response.candidates and hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason: reason = response.prompt_feedback.block_reason.name; logging.error(f"Gemini prompt blocked: {reason}."); return {"error": f"Prompt blocked ({reason})."}
            if not response.candidates or not hasattr(response.candidates[0].content, 'parts') or not response.candidates[0].content.parts[0].text: logging.error("Gemini response empty/malformed."); return {"error": "AI returned empty response."}

            generated_text_raw = response.text
            
        # Process the text to ensure proper citation formatting
        generated_text_processed = preprocess_llm_text_for_citations(generated_text_raw)

        # Convert markdown to HTML
        markdown_parser = mistune.create_markdown(renderer='html', plugins=['strikethrough', 'footnotes', 'table', 'task_lists'])
        html_answer = markdown_parser(generated_text_processed)
        
        # Process citation markers in HTML with improved regex
        # This regex carefully matches individual citation markers
        html_answer = re.sub(
            r"(?<![!\]/a-zA-Z0-9])\[(\d+)\](?!\])", 
            lambda m: f"<sup><a href='#' class='citation-marker' data-citation-index='{int(m.group(1))-1}' aria-label='Citation {m.group(1)}'>[{m.group(1)}]</a></sup> ",  # Note the space after </sup>
            html_answer
        )
        
        # Clean up any excessive spaces that might have been added
        html_answer = re.sub(r"(\s{2,})", " ", html_answer)
        
        logging.info(f"Successfully generated response. Processed length: {len(generated_text_processed)}")
        return {"answer_raw": generated_text_processed, "answer_html": html_answer}

    except Exception as e:
        # Error handling (unchanged)
        logging.error(f"Error calling Gemini API: {e}", exc_info=True)
        err_msg = f"AI communication error: {str(e)}"

        if "API key not valid" in str(e) or "invalidated api key" in str(e).lower():
            err_msg = "Invalid Gemini API Key."
        elif "Rate limit exceeded" in str(e) or "quota" in str(e).lower() or "429" in str(e):
            err_msg = "AI model rate limit exceeded. Please try again later or check your API quota."
        elif "User location is not supported" in str(e):
            err_msg = "AI model access restricted by location."
        elif "exceeds the maximum token limit" in str(e).lower() or "token limit" in str(e).lower():
            err_msg = "Content too large for AI model. Try with fewer sources or 'Quick' mode."

        return {"error": err_msg}

# --- UPDATED generate_docx ---
def generate_docx(report_data: dict) -> BytesIO:
    """Generates DOCX, fixing content addition and conditional source list."""
    document = Document()
    try:
        default_font = document.styles['Normal'].font; default_font.name = 'Arial'; default_font.size = Pt(11)
        query = report_data.get('query', 'Untitled Report')
        document.add_heading("Research Report", level=1)
        p = document.add_paragraph(); p.add_run("Query: ").bold = True; p.add_run(query); document.add_paragraph()

        answer_content_processed = report_data.get('answer_raw', 'Content not available.')
        # --- FIX: Refined Markdown to DOCX paragraph handling ---
        current_paragraph = None
        for line in answer_content_processed.split('\n'):
            stripped_line = line.strip()

            # Heading detection
            if stripped_line.startswith('#'):
                current_paragraph = None # End previous paragraph before heading
                level = stripped_line.count('#', 0, 3)
                heading_text = stripped_line.lstrip('# ').strip()
                if heading_text: document.add_heading(heading_text, level=min(level, 3))
                continue

            # List detection (basic, no nesting support here)
            is_list_item = False
            list_content = None
            list_style = 'List Bullet' # Default
            if stripped_line.startswith(('* ', '- ')):
                is_list_item = True
                list_content = stripped_line[2:]
                list_style = 'List Bullet'
            elif re.match(r"^\d+\.\s+", stripped_line):
                 is_list_item = True
                 list_content = re.sub(r"^\d+\.\s+", "", stripped_line)
                 list_style = 'List Number'

            if is_list_item:
                current_paragraph = None # End previous paragraph before list item
                if list_content: # Avoid adding empty list items
                    document.add_paragraph(list_content, style=list_style)
                continue # Move to next line after handling list item

            # Normal paragraph text
            if stripped_line:
                if current_paragraph is None:
                    # Start a new paragraph if not continuing one or after heading/list
                    current_paragraph = document.add_paragraph(stripped_line)
                else:
                    # Add line break *within* the current paragraph
                    current_paragraph.add_run("\n" + stripped_line) # Use Word line break
            elif current_paragraph is not None:
                 # Empty line signifies paragraph break
                 current_paragraph = None

        # --- End FIX ---

        # --- Conditional Sources Section ---
        document.add_page_break()
        research_depth = report_data.get('research_depth', 'quick') # Correctly get depth
        sources_heading_text = "References" if research_depth == 'deep' else "Sources Cited"
        document.add_heading(sources_heading_text, level=1)
        sources = report_data.get('sources', [])
        if sources:
            for i, source in enumerate(sources):
                p = document.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.first_line_indent = Inches(-0.25)
                p.add_run(f"[{i+1}] ").bold = True; title = source.get('title', 'Source Title Unavailable'); url = source.get('url', '')
                p.add_run(f"{title}")
                if research_depth == 'deep' and url: p.add_run(f" ({url})") # URL only for deep
        else: document.add_paragraph("No sources were cited.")

        file_stream = BytesIO(); document.save(file_stream); file_stream.seek(0)
        logging.info(f"DOCX generated successfully (Depth: {research_depth})")
        return file_stream
    except Exception as e: logging.error(f"Error generating DOCX: {e}", exc_info=True); raise


# --- FIXED generate_pdf ---
def generate_pdf(report_data: dict) -> BytesIO | None:
    """Generates PDF, ensuring depth check and using updated CSS."""
    try:
        query = report_data.get('query', 'Untitled Report')
        answer_html_content = report_data.get('answer_html', '<p>Content not available.</p>')
        sources = report_data.get('sources', [])
        research_depth = report_data.get('research_depth', 'quick') # Get depth

        source_list_html = ""
        if sources:
            list_class = "sources-list " + ("deep-list" if research_depth == 'deep' else "quick-list")
            source_list_html = f'<ul class="{list_class}">'
            for i, source in enumerate(sources):
                title = html.escape(source.get('title', 'Source Title Unavailable'))
                url = html.escape(source.get('url', ''))
                preview = html.escape(source.get('text_preview', ''))
                source_list_html += f'<li><span class="source-number">[{i+1}]</span> <span class="source-title">{title}</span>'
                if research_depth == 'deep' and url: source_list_html += f' <a href="{url}" class="source-url">({url})</a>'
                elif research_depth == 'quick' and url: source_list_html += f'<br><a href="{url}" class="source-url">{url}</a>'
                if research_depth == 'quick' and preview: source_list_html += f'<p class="source-preview">{preview}...</p>'
                source_list_html += '</li>'
            source_list_html += "</ul>"
        else: source_list_html = "<p>No sources were cited.</p>"

        sources_heading_text = "References" if research_depth == 'deep' else "Sources Cited"

        # --- PDF HTML and CSS ---
        # In the generate_pdf function, update the CSS for citation markers:
        # In the generate_pdf function, update the CSS for citation markers:
        pdf_html = f"""
        <!DOCTYPE html><html><head><meta charset="UTF-8"><title>Report: {html.escape(query)}</title><style>
            /* Base styles */
            @page {{ size: a4 portrait; margin: 2cm 1.5cm; @frame header_frame {{ -pdf-frame-content: header_content; left: 1.5cm; width: 18cm; top: 1cm; height: 1cm; }} @frame footer_frame {{ -pdf-frame-content: footer_content; left: 1.5cm; width: 18cm; top: 26.7cm; height: 1cm; }} }}
            body {{ font-family: Arial, Helvetica, sans-serif; font-size: 10pt; line-height: 1.4; color: #333; }}
            #header_content, #footer_content {{ font-size: 9pt; color: #777; }} #header_content {{ text-align: left; }} #footer_content {{ text-align: right; }}
            h1.report-title {{ font-size: 16pt; color: #2c3e50; margin-bottom: 5px; font-weight: bold; }}
            h2.query-title {{ font-size: 11pt; color: #555; font-weight: normal; margin-bottom: 20px; border-bottom: 1px solid #eee; padding-bottom: 10px; }}
            /* Content Styles */
            .answer-content h2 {{ font-size: 14pt; color: #2c3e50; margin-top: 1.2em; margin-bottom: 0.6em; font-weight: bold; border-bottom: 1px solid #ccc; padding-bottom: 3px; }}
            .answer-content h3 {{ font-size: 12pt; color: #2c3e50; margin-top: 1em; margin-bottom: 0.5em; font-weight: bold; border-bottom: 1px solid #eee; padding-bottom: 2px; }}
            .answer-content p {{ margin-bottom: 0.8em; text-align: justify; }}
            .answer-content ul, .answer-content ol {{ margin-left: 20px; margin-bottom: 0.8em; }} .answer-content ul {{ list-style-type: disc; }} .answer-content ol {{ list-style-type: decimal; }} .answer-content li {{ margin-bottom: 0.3em; }}
            .answer-content a {{ color: #0066cc; text-decoration: underline; }}
            /* Citation Marker Style - Improved for PDF */
            .answer-content sup {{ display: inline-block; margin: 0 1px; }}
            sup a.citation-marker {{
                font-size: 0.8em;
                vertical-align: super;
                padding: 1px 2px;
                margin: 0 1px;
                color: #0066cc;
                text-decoration: none;
                background-color: #f0f0f0;
                border-radius: 2px;
                display: inline-block;
            }}
            /* Source List Styles */
            h3.sources-heading {{ font-size: 14pt; color: #2c3e50; margin-top: 25px; margin-bottom: 10px; border-bottom: 1px solid #ccc; padding-bottom: 5px; page-break-before: always; }}
            ul.sources-list {{ list-style-type: none; padding-left: 5px; margin-top: 0; }}
            ul.sources-list li {{ margin-bottom: 8px; font-size: 9pt; line-height: 1.3; }}
            span.source-number {{ font-weight: bold; margin-right: 5px; color: #111; }} span.source-title {{ color: #333; font-weight: 600; }}
            a.source-url {{ color: #0066cc; text-decoration: none; font-size: 0.9em; word-break: break-all; }}
            ul.quick-list li {{ background-color: #f8f8f8; padding: 8px; border: 1px solid #eee; border-radius: 3px;}}
            ul.quick-list a.source-url {{ display: block; margin-top: 2px; }}
            p.source-preview {{ color: #555; font-size: 0.85em; margin-top: 5px; margin-bottom: 0; padding-left: 15px; border-left: 2px solid #ddd; max-height: 4.5em; overflow: hidden; }}
            ul.deep-list li {{ background-color: transparent; padding: 2px 0; border: none; }}
            ul.deep-list a.source-url {{ display: inline; margin-left: 5px; }}
            ul.deep-list p.source-preview {{ display: none; }}
            </style></head><body>
            <div id="header_content">Research Report</div><div id="footer_content">Page <pdf:pagenumber> of <pdf:pagecount></div>
            <h1 class="report-title">Research Report</h1><h2 class="query-title">Query: {html.escape(query)}</h2>
            <div class="answer-content">{answer_html_content}</div>
            <h3 class="sources-heading">{sources_heading_text}</h3>
            {source_list_html}
            </body></html>"""

        pdf_stream = BytesIO(); pisa_status = pisa.CreatePDF(src=pdf_html, dest=pdf_stream)
        if pisa_status.err: logging.error(f"PDF generation failed: {pisa_status.err}"); return None
        else: pdf_stream.seek(0); logging.info(f"PDF generated successfully (Depth: {research_depth})"); return pdf_stream
    except Exception as e: logging.error(f"Exception during PDF generation: {e}", exc_info=True); return None


# --- Flask App Initialization and Routes ---
def create_app():
    app = Flask(__name__)
    # (Keep app config, index route)
    app.config['SECRET_KEY'] = os.getenv("FLASK_SECRET_KEY", "default-fallback-secret-key")
    app.config['GEMINI_API_KEY'] = os.getenv("GEMINI_API_KEY")
    if not app.config['GEMINI_API_KEY']: logging.warning("GEMINI_API_KEY not found.")
    else: logging.info("GEMINI_API_KEY found.")
    @app.route('/')
    def index(): return render_template('index.html')

    @app.route('/process', methods=['POST'])
    def process_query_route():
        # (Keep route logic as before - stores depth correctly)
        start_time = time.time(); data = request.get_json(); query = data.get('query', '').strip()
        research_depth = data.get('depth', 'quick')
        if not request.is_json: return jsonify({"error": "Request must be JSON"}), 400
        if not query: return jsonify({"error": "Query cannot be empty."}), 400
        api_key = current_app.config.get('GEMINI_API_KEY')
        if not api_key: return jsonify({"error": "AI model not configured."}), 500
        logging.info(f"Processing query: '{query}' [Depth: {research_depth}]")
        try:
            num_results = DEEP_SEARCH_RESULTS if research_depth == 'deep' else QUICK_SEARCH_RESULTS
            search_results = perform_search(query, num_results=num_results)
            urls_to_scrape = [r['url'] for r in search_results if r.get('url')]
            url_to_title_map = {r['url']: r.get('title') for r in search_results if r.get('url')}
            if not urls_to_scrape: return jsonify({"error": "Could not find relevant web sources."}), 404
            logging.info(f"Attempting to scrape {len(urls_to_scrape)} URLs.")
            scraped_data = scrape_urls(urls_to_scrape)
            if not scraped_data: return jsonify({"error": "Failed to retrieve usable content from web sources."}), 500
            success_rate = len(scraped_data) / len(urls_to_scrape) if urls_to_scrape else 0
            if success_rate < 0.4: logging.warning(f"Low scrape success rate ({success_rate:.1%}).")
            synthesis_result = synthesize_with_gemini(query, scraped_data, api_key, research_depth)
            sources_final = []
            for item in scraped_data: preview = item.get('text', '')[:SOURCE_PREVIEW_LENGTH]; sources_final.append({"id": item["id"], "url": item["url"], "title": url_to_title_map.get(item["url"], f"Source {item['id']+1}"), "text_preview": preview})
            if not synthesis_result or "error" in synthesis_result: error_msg = synthesis_result.get("error", "AI synthesis failed.") if synthesis_result else "AI synthesis failed."; logging.error(f"Synthesis failed: {error_msg}"); return jsonify({"error": f"AI Synthesis Error: {error_msg}", "sources": sources_final, "research_depth": research_depth}), 500
            report_id = str(uuid.uuid4())
            # Ensure depth is stored correctly
            report_full_data = {"query": query, "answer_raw": synthesis_result["answer_raw"], "answer_html": synthesis_result["answer_html"], "sources": sources_final, "research_depth": research_depth}
            add_to_report_store(report_id, report_full_data)
            end_time = time.time(); logging.info(f"Success query '{query}' in {end_time - start_time:.2f}s. ID: {report_id}")
            return jsonify({"answer_html": synthesis_result["answer_html"], "sources": sources_final, "report_id": report_id, "research_depth": research_depth}) # Return depth
        except Exception as e: logging.exception(f"Unexpected error in '/process': {e}"); return jsonify({"error": "An unexpected internal server error."}), 500

    @app.route('/download/docx/<report_id>')
    def download_docx_route(report_id):
        # (generate_docx now correctly checks depth from report_data)
        report_data = report_store.get(report_id)
        if not report_data: return "Report not found or has expired.", 404
        try: file_stream = generate_docx(report_data); query_slug = "".join(c if c.isalnum() else "_" for c in report_data.get('query', 'report'))[:30].strip('_'); filename = f"Research_Report_{query_slug or 'report'}.docx"; logging.info(f"Serving DOCX: {filename}"); return send_file(file_stream, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename)
        except Exception as e: logging.error(f"Error generating/sending DOCX: {e}", exc_info=True); return "Error generating DOCX file.", 500 # Log exception detail

    @app.route('/download/pdf/<report_id>')
    def download_pdf_route(report_id):
        # (generate_pdf now correctly checks depth from report_data)
        report_data = report_store.get(report_id)
        if not report_data: return "Report not found or has expired.", 404
        try:
            pdf_stream = generate_pdf(report_data)
            if not pdf_stream: logging.error(f"PDF generation returned None for report {report_id}"); return "Error generating PDF file (check server logs).", 500
            query_slug = "".join(c if c.isalnum() else "_" for c in report_data.get('query', 'report'))[:30].strip('_'); filename = f"Research_Report_{query_slug or 'report'}.pdf"; logging.info(f"Serving PDF: {filename}"); response = make_response(pdf_stream.getvalue()); response.headers['Content-Type'] = 'application/pdf'; response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'; return response
        except Exception as e: logging.error(f"Error generating/sending PDF: {e}", exc_info=True); return "Error generating PDF file.", 500 # Log exception detail

    return app

# --- Main Execution ---
if __name__ == '__main__':
    app = create_app()
    app.run(debug=True, host='0.0.0.0', port=5000)

# --- END OF FILE app.py ---
